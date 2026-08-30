import os
from fpdf import FPDF
from docx import Document
from tools.base import Tool
from tools.file_tools import is_safe_path

class ExportPDFTool(Tool):
    name = "export_pdf"
    def execute(self, path: str, content: str, title: str = "Report", **kwargs) -> str:
        if not is_safe_path(path): return "Error: Access denied."
        try:
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            pdf.cell(200, 10, txt=title, ln=1, align='C')
            # Handle unicode gracefully by replacing unsupported chars
            pdf.multi_cell(0, 10, txt=content.encode('latin-1', 'replace').decode('latin-1'))
            pdf.output(path)
            return f"PDF saved to {path}"
        except Exception as e: return f"PDF Export Error: {str(e)}"

class ExportDocxTool(Tool):
    name = "export_docx"
    def execute(self, path: str, content: str, title: str = "Report", **kwargs) -> str:
        if not is_safe_path(path): return "Error: Access denied."
        try:
            doc = Document()
            doc.add_heading(title, 0)
            doc.add_paragraph(content)
            doc.save(path)
            return f"DOCX saved to {path}"
        except Exception as e: return f"DOCX Export Error: {str(e)}"

class ExportMarkdownTool(Tool):
    name = "export_markdown"
    def execute(self, path: str, content: str, **kwargs) -> str:
        if not is_safe_path(path): return "Error: Access denied."
        try:
            with open(path, 'w') as f:
                f.write(content)
            return f"Markdown saved to {path}"
        except Exception as e: return f"MD Export Error: {str(e)}"
